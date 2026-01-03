# run_all.py
# =========================================================
# 两套任务 + 两套评估（KFold & GroupKFold）+ 自动生成表格
#
# 任务A：X=[钻速(mm/s), 角速度(1/s), 推力(kN), 扭矩(N·m), 岩石可钻性指数] -> Y=[单轴抗压强度(MPa)]
# 任务B：X=[钻速(mm/s), 角速度(1/s), 推力(kN), 扭矩(N·m)] -> Y=[岩石可钻性指数, 单轴抗压强度(MPa)]
#
# 模型：MLP / Random_Forest / SVR / XGBoost（若未安装xgboost则自动跳过）
# 指标：R2 / RMSE / MAE（mean ± std）
#
# 评估：
#  - KFold(5) : 常规5折
#  - GroupKFold : 按 UCS 分组
#
# 输出：
#  - results_runall/cv_results_long.csv  （长表：每折结果聚合 mean/std）
#  - results_runall/summary_tables/*.csv （对比总表）
#  - results_runall/latex_tables/*.tex   （论文LaTeX表格）
#  - results_runall/word_tables/*.docx   （论文Word表格）
# =========================================================
'''
这一版为将所有模型的构建全部放在主程序中，没有做模块化处理，保留初版代码
'''
import os
from datetime import datetime
import warnings
from dataclasses import dataclass
from typing import Dict, List, Tuple, Optional

import numpy as np
import pandas as pd

from sklearn.model_selection import KFold, GroupKFold
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import StandardScaler
from sklearn.metrics import r2_score, mean_squared_error, mean_absolute_error
from sklearn.multioutput import MultiOutputRegressor

from sklearn.neural_network import MLPRegressor
from sklearn.ensemble import RandomForestRegressor
from sklearn.svm import SVR

warnings.filterwarnings("ignore")

# -------------------------
# 1) 严格列名
# -------------------------
COL_V   = "钻速(mm/s)"
COL_W   = "角速度(1/s)"
COL_F   = "推力(kN)"
COL_TQ  = "扭矩(N·m)"
COL_ID  = "岩石可钻性指数"
COL_UCS = "单轴抗压强度(MPa)"

TASK_A_X = [COL_V, COL_W, COL_F, COL_TQ, COL_ID]
TASK_A_Y = [COL_UCS]

TASK_B_X = [COL_V, COL_W, COL_F, COL_TQ]
TASK_B_Y = [COL_ID, COL_UCS]

# -------------------------
# 2) 路径配置
# -------------------------
ROOT_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_PATH = os.path.join(ROOT_DIR, "data", "raw", "cz_drill.xlsx")
RUN_TAG = datetime.now().strftime("%Y.%m.%d.%H.%M.%S")
OUT_DIR = os.path.join(ROOT_DIR, "results", RUN_TAG)


# -------------------------
# 3) 基础函数
# -------------------------
def ensure_dir(p: str) -> None:
    os.makedirs(p, exist_ok=True)

def rmse(y_true, y_pred) -> float:
    return float(np.sqrt(mean_squared_error(y_true, y_pred)))

def eval_metrics_1d(y_true: np.ndarray, y_pred: np.ndarray) -> Dict[str, float]:
    return {
        "R2": float(r2_score(y_true, y_pred)),
        "RMSE": rmse(y_true, y_pred),
        "MAE": float(mean_absolute_error(y_true, y_pred)),
    }

def to_2d(arr) -> np.ndarray:
    arr = np.asarray(arr)
    if arr.ndim == 1:
        return arr.reshape(-1, 1)
    return arr

def format_mean_std(mean: float, std: float, nd: int = 3) -> str:
    return f"{mean:.{nd}f}±{std:.{nd}f}"

def assert_columns(df: pd.DataFrame, cols: List[str], tag: str):
    missing = [c for c in cols if c not in df.columns]
    if missing:
        raise KeyError(
            f"[{tag}] 缺少列：{missing}\n"
            f"当前表头：{list(df.columns)}\n"
            f"请确保Excel表头与脚本列名完全一致（含大小写、单位、·符号）。"
        )

def load_data(path: str) -> pd.DataFrame:
    if not os.path.exists(path):
        raise FileNotFoundError(f"找不到数据文件：{path}")
    if path.lower().endswith((".xlsx", ".xls")):
        df = pd.read_excel(path)
    else:
        df = pd.read_csv(path)
    df = df.dropna(how="all").copy()
    return df

def clean_xy(df: pd.DataFrame, x_cols: List[str], y_cols: List[str]) -> Tuple[pd.DataFrame, pd.DataFrame]:
    sub = df[x_cols + y_cols].copy()
    for c in x_cols + y_cols:
        sub[c] = pd.to_numeric(sub[c], errors="coerce")
    sub = sub.dropna(axis=0, how="any").reset_index(drop=True)
    return sub[x_cols], sub[y_cols]

# -------------------------
# 4) 模型构建
# -------------------------
def build_models(random_state: int = 42) -> Dict[str, object]:
    models: Dict[str, object] = {}

    models["MLP"] = Pipeline([
        ("scaler", StandardScaler()),
        ("reg", MLPRegressor(
            hidden_layer_sizes=(64, 32),
            activation="relu",
            solver="adam",
            alpha=1e-4,
            learning_rate_init=1e-3,
            max_iter=3000,
            random_state=random_state,
            early_stopping=True,
            n_iter_no_change=40
        ))
    ])

    models["SVR"] = Pipeline([
        ("scaler", StandardScaler()),
        ("reg", SVR(C=50, epsilon=0.05, kernel="rbf"))
    ])

    models["Random_Forest"] = RandomForestRegressor(
        n_estimators=800,
        random_state=random_state,
        n_jobs=-1
    )

    # XGBoost 可选
    try:
        from xgboost import XGBRegressor
        models["XGBoost"] = XGBRegressor(
            n_estimators=1200,
            learning_rate=0.03,
            max_depth=6,
            subsample=0.9,
            colsample_bytree=0.9,
            reg_lambda=1.0,
            random_state=random_state,
            n_jobs=-1
        )
    except Exception:
        print("[WARN] xgboost 未安装或不可用：将跳过 XGBoost。")

    return models

def wrap_multioutput(model, n_targets: int):
    if n_targets <= 1:
        return model
    # RF 原生支持多输出；其他用 MultiOutputRegressor
    if isinstance(model, RandomForestRegressor):
        return model
    return MultiOutputRegressor(model)

# -------------------------
# 5) CV 评估核心（支持KFold & GroupKFold）
# -------------------------
@dataclass
class AggRow:
    eval_scheme: str
    task: str
    model: str
    target: str
    metric: str
    mean: float
    std: float

def run_cv(
    X: pd.DataFrame,
    Y: pd.DataFrame,
    models: Dict[str, object],
    task_name: str,
    eval_scheme: str,
    splitter,
    groups: Optional[np.ndarray] = None
) -> List[AggRow]:

    Xv = X.values
    Yv = to_2d(Y.values)
    target_names = list(Y.columns)
    n_targets = Yv.shape[1]

    out_rows: List[AggRow] = []

    for model_name, base_model in models.items():
        # 收集每折每目标每指标
        fold_store = {t: {"R2": [], "RMSE": [], "MAE": []} for t in target_names}

        if groups is None:
            split_iter = splitter.split(Xv, Yv)
        else:
            split_iter = splitter.split(Xv, Yv, groups)

        for tr_idx, te_idx in split_iter:
            Xtr, Xte = Xv[tr_idx], Xv[te_idx]
            Ytr, Yte = Yv[tr_idx], Yv[te_idx]

            model = wrap_multioutput(base_model, n_targets)
            model.fit(Xtr, Ytr)
            Ypred = to_2d(model.predict(Xte))

            for j, tname in enumerate(target_names):
                m = eval_metrics_1d(Yte[:, j], Ypred[:, j])
                for k, v in m.items():
                    fold_store[tname][k].append(v)

        # 聚合：每目标
        for tname in target_names:
            for metric in ["R2", "RMSE", "MAE"]:
                vals = np.array(fold_store[tname][metric], dtype=float)
                out_rows.append(AggRow(
                    eval_scheme=eval_scheme,
                    task=task_name,
                    model=model_name,
                    target=tname,
                    metric=metric,
                    mean=float(vals.mean()),
                    std=float(vals.std(ddof=1)) if len(vals) > 1 else 0.0
                ))

        # 任务B的 macro avg
        if n_targets > 1:
            for metric in ["R2", "RMSE", "MAE"]:
                per_target_means = np.array(
                    [np.mean(fold_store[t][metric]) for t in target_names],
                    dtype=float
                )
                out_rows.append(AggRow(
                    eval_scheme=eval_scheme,
                    task=task_name,
                    model=model_name,
                    target="MACRO_AVG",
                    metric=metric,
                    mean=float(per_target_means.mean()),
                    std=float(per_target_means.std(ddof=1)) if len(per_target_means) > 1 else 0.0
                ))

    return out_rows

# -------------------------
# 6) 生成“模型×指标”总表（mean±std） + 导出 LaTeX/Word
# -------------------------
def build_summary_table(df_long: pd.DataFrame, task: str, eval_scheme: str, target: str) -> pd.DataFrame:
    """
    返回：index=model, columns=[R2,RMSE,MAE] 的表，值为 "mean±std"
    """
    sub = df_long[(df_long["task"] == task) & (df_long["eval_scheme"] == eval_scheme) & (df_long["target"] == target)].copy()
    if sub.empty:
        return pd.DataFrame()

    sub["mean_std"] = sub.apply(lambda r: format_mean_std(r["mean"], r["std"], nd=3), axis=1)
    pivot = sub.pivot_table(index="model", columns="metric", values="mean_std", aggfunc="first")
    pivot = pivot[["R2", "RMSE", "MAE"]]  # 固定列顺序
    pivot = pivot.reset_index()
    return pivot

def save_latex_table(df: pd.DataFrame, out_path: str, caption: str, label: str):
    if df.empty:
        return
    latex = df.to_latex(index=False, escape=False, caption=caption, label=label)
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(latex)

def save_word_tables(all_tables: List[Tuple[str, pd.DataFrame]], out_docx: str):
    """
    all_tables: [(标题, DataFrame), ...]
    """
    try:
        from docx import Document
    except Exception:
        print("[WARN] python-docx 未安装或不可用：将跳过 Word 表格输出。")
        return

    doc = Document()
    doc.add_heading("模型评估结果汇总表（mean±std）", level=1)

    for title, df in all_tables:
        if df.empty:
            continue
        doc.add_heading(title, level=2)

        table = doc.add_table(rows=1, cols=df.shape[1])
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr_cells[i].text = str(col)

        for _, row in df.iterrows():
            cells = table.add_row().cells
            for i, col in enumerate(df.columns):
                cells[i].text = str(row[col])

        doc.add_paragraph("")  # 空行

    doc.save(out_docx)

# -------------------------
# 7) 主程序
# -------------------------
def main():
    ensure_dir(OUT_DIR)
    ensure_dir(os.path.join(OUT_DIR, "summary_tables"))
    ensure_dir(os.path.join(OUT_DIR, "latex_tables"))
    ensure_dir(os.path.join(OUT_DIR, "word_tables"))

    df = load_data(DATA_PATH)

    # 检查列
    assert_columns(df, TASK_A_X + TASK_A_Y, "Task_A")
    assert_columns(df, TASK_B_X + TASK_B_Y, "Task_B")

    models = build_models(random_state=42)

    # -------- Task A 数据
    Xa, Ya = clean_xy(df, TASK_A_X, TASK_A_Y)

    # -------- Task B 数据
    Xb, Yb = clean_xy(df, TASK_B_X, TASK_B_Y)

    # -------- 两套评估
    # 1) KFold(5)
    kf = KFold(n_splits=5, shuffle=True, random_state=42)

    # 2) GroupKFold：按 UCS 分组（跨工况泛化）
    #    groups 使用 UCS 的“离散组”标签
    groups_a = Ya[COL_UCS].values
    groups_b = Yb[COL_UCS].values

    # GroupKFold 的 n_splits 不能超过 group 数量
    n_groups = len(np.unique(groups_a))
    gkf = GroupKFold(n_splits=min(4, n_groups))

    # -------- 跑评估并汇总成长表
    rows: List[AggRow] = []
    rows += run_cv(Xa, Ya, models, task_name="Task_A: 5->UCS", eval_scheme="KFold-5", splitter=kf, groups=None)
    rows += run_cv(Xa, Ya, models, task_name="Task_A: 5->UCS", eval_scheme="GroupKFold-UCS", splitter=gkf, groups=groups_a)

    rows += run_cv(Xb, Yb, models, task_name="Task_B: 4->[Id,UCS]", eval_scheme="KFold-5", splitter=kf, groups=None)
    rows += run_cv(Xb, Yb, models, task_name="Task_B: 4->[Id,UCS]", eval_scheme="GroupKFold-UCS", splitter=gkf, groups=groups_b)

    df_long = pd.DataFrame([r.__dict__ for r in rows])
    long_csv = os.path.join(OUT_DIR, "cv_results_long.csv")
    df_long.to_csv(long_csv, index=False, encoding="utf-8-sig")

    # -------- 生成对比总表（模型×指标）
    # 表格清单：每个（任务×评估×目标）一张表
    tables_for_word: List[Tuple[str, pd.DataFrame]] = []

    combos = [
        ("Task_A: 5->UCS", "KFold-5", COL_UCS),
        ("Task_A: 5->UCS", "GroupKFold-UCS", COL_UCS),

        ("Task_B: 4->[Id,UCS]", "KFold-5", COL_ID),
        ("Task_B: 4->[Id,UCS]", "KFold-5", COL_UCS),
        ("Task_B: 4->[Id,UCS]", "KFold-5", "MACRO_AVG"),

        ("Task_B: 4->[Id,UCS]", "GroupKFold-UCS", COL_ID),
        ("Task_B: 4->[Id,UCS]", "GroupKFold-UCS", COL_UCS),
        ("Task_B: 4->[Id,UCS]", "GroupKFold-UCS", "MACRO_AVG"),
    ]

    for task, scheme, target in combos:
        tab = build_summary_table(df_long, task=task, eval_scheme=scheme, target=target)
        if tab.empty:
            continue

        # 保存 CSV
        safe_name = f"{task.replace(':','_').replace('->','to').replace('[','').replace(']','').replace(',','_').replace(' ','')}_{scheme}_{target}.csv"
        csv_path = os.path.join(OUT_DIR, "summary_tables", safe_name)
        tab.to_csv(csv_path, index=False, encoding="utf-8-sig")

        # 保存 LaTeX
        caption = f"{task}（{scheme}），目标：{target} 的模型性能对比（mean±std）"
        label = f"tab:{task.replace(' ','_').replace(':','_')}_{scheme}_{target}"
        tex_name = safe_name.replace(".csv", ".tex")
        tex_path = os.path.join(OUT_DIR, "latex_tables", tex_name)
        save_latex_table(tab, tex_path, caption=caption, label=label)

        # Word 收集
        title = f"{task} | {scheme} | 目标：{target}"
        tables_for_word.append((title, tab))

    # -------- 输出 Word 文档（把所有表放一份docx里）
    word_path = os.path.join(OUT_DIR, "word_tables", "paper_tables.docx")
    save_word_tables(tables_for_word, word_path)

    # -------- 终端提示
    print("\n==================== 完成 ====================")
    print(f"数据文件：{DATA_PATH}")
    print(f"长表输出：{long_csv}")
    print(f"对比总表：{os.path.join(OUT_DIR, 'summary_tables')}")
    print(f"LaTeX表格：{os.path.join(OUT_DIR, 'latex_tables')}")
    print(f"Word表格：{word_path}")
    print("================================================\n")

if __name__ == "__main__":
    main()