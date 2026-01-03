import pandas as pd
from typing import List, Tuple


def format_mean_std(mean: float, std: float, nd: int = 3) -> str:
    if pd.isna(mean) or pd.isna(std):
        return ""
    return f"{mean:.{nd}f}±{std:.{nd}f}"


def build_summary_table(df_long: pd.DataFrame, task: str, eval_scheme: str, target: str) -> pd.DataFrame:
    """
    返回：columns = [model, R2, RMSE, MAE]，每格是 "mean±std"
    """
    sub = df_long[
        (df_long["task"] == task) &
        (df_long["eval_scheme"] == eval_scheme) &
        (df_long["target"] == target)
    ].copy()

    if sub.empty:
        return pd.DataFrame()

    sub["mean_std"] = sub.apply(lambda r: format_mean_std(r["mean"], r["std"], nd=3), axis=1)

    pivot = sub.pivot_table(index="model", columns="metric", values="mean_std", aggfunc="first")

    # 固定列顺序 + 缺失列兜底
    order = ["R2", "RMSE", "MAE"]
    for m in order:
        if m not in pivot.columns:
            pivot[m] = ""
    pivot = pivot[order].reset_index()

    return pivot


def build_all_tables(df_long: pd.DataFrame) -> List[Tuple[str, pd.DataFrame]]:
    """
    批量生成所有组合的汇总表，返回 [(title, df), ...]
    title 适合直接当 Word 二级标题/LaTeX caption 的基础文本
    """
    tables: List[Tuple[str, pd.DataFrame]] = []

    tasks = df_long["task"].dropna().unique().tolist()
    schemes = df_long["eval_scheme"].dropna().unique().tolist()

    for task in tasks:
        for scheme in schemes:
            # 每个组合下有哪些 target
            targets = (
                df_long[(df_long["task"] == task) & (df_long["eval_scheme"] == scheme)]["target"]
                .dropna().unique().tolist()
            )
            for target in targets:
                df = build_summary_table(df_long, task, scheme, target)
                title = f"{task} | {scheme} | target={target}"
                tables.append((title, df))

    return tables


def save_latex_table(df: pd.DataFrame, out_path: str, caption: str, label: str):
    if df.empty:
        return
    latex = df.to_latex(index=False, escape=False, caption=caption, label=label)
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(latex)


def save_word_tables(all_tables: List[Tuple[str, pd.DataFrame]], out_docx: str):
    """
    all_tables: [(标题, DataFrame), ...]
    """
    try:
        from docx import Document
    except Exception:
        print("[WARN] python-docx 未安装或不可用：将跳过 Word 表格输出。")
        return

    doc = Document()
    doc.add_heading("模型评估结果汇总表（mean±std）", level=1)

    for title, df in all_tables:
        if df.empty:
            continue
        doc.add_heading(title, level=2)

        table = doc.add_table(rows=1, cols=df.shape[1])
        table.style = "Table Grid"

        hdr_cells = table.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr_cells[i].text = str(col)

        for _, row in df.iterrows():
            cells = table.add_row().cells
            for i, col in enumerate(df.columns):
                cells[i].text = str(row[col])

        doc.add_paragraph("")

    doc.save(out_docx)
